import io
from fastapi import HTTPException, status
import pypdf
import docx

# Minimum characters required from pypdf extraction before OCR fallback is triggered.
# Scanned PDFs typically return 0-50 chars of garbage or nothing at all.
MIN_TEXT_THRESHOLD = 100


def _ocr_pdf_bytes(stream_bytes: bytes) -> str:
    """
    OCR Fallback: Rasterize each PDF page using pdf2image (Poppler) and run
    Tesseract OCR over each page image.

    - Zero disk writes: all processing happens in RAM via BytesIO.
    - Gracefully degrades to empty string if OCR dependencies are missing.
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract

        images = convert_from_bytes(stream_bytes, dpi=250)
        pages_text = []
        for img in images:
            page_text = pytesseract.image_to_string(img, lang="eng")
            if page_text.strip():
                pages_text.append(page_text.strip())
        return "\n\n".join(pages_text)
    except ImportError:
        # OCR libraries not available — silent degradation
        return ""
    except Exception:
        # Tesseract not installed or rasterization failure — silent degradation
        return ""


def extract_text_from_pdf(stream_bytes: bytes) -> tuple[str, bool]:
    """
    Extract raw text from PDF bytes in memory.

    Primary path:   pypdf embedded-text extraction (fast, lossless).
    Fallback path:  Tesseract OCR via pdf2image (for scanned/image PDFs).

    Returns:
        (text, ocr_used) — text is the extracted content; ocr_used is True
        when the OCR fallback was triggered.
    """
    try:
        reader = pypdf.PdfReader(io.BytesIO(stream_bytes))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        extracted = "\n\n".join(pages_text)

        if len(extracted.strip()) < MIN_TEXT_THRESHOLD:
            # Likely a scanned / image-only PDF — trigger OCR fallback
            ocr_text = _ocr_pdf_bytes(stream_bytes)
            if ocr_text.strip():
                return ocr_text, True
            # OCR also returned nothing (truly empty or corrupt) — return best available
            return extracted, False

        return extracted, False

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to extract text from PDF file: {str(e)}"
        )


def extract_text_from_docx(stream_bytes: bytes) -> str:
    """Extract raw text from DOCX bytes in memory."""
    try:
        doc = docx.Document(io.BytesIO(stream_bytes))
        paragraphs_text = []

        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs_text.append(p.text.strip())

        # Extract table text if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs_text.append(row_text)

        return "\n\n".join(paragraphs_text)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to extract text from DOCX file: {str(e)}"
        )


def extract_text_from_txt(stream_bytes: bytes) -> str:
    """Extract raw text from TXT / MD bytes in memory."""
    try:
        return stream_bytes.decode('utf-8', errors='replace')
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to decode text file: {str(e)}"
        )


def extract_document_text(filename: str, content_type: str | None, stream_bytes: bytes) -> tuple[str, bool]:
    """
    Route document to proper in-memory extractor based on extension and content-type.

    Returns:
        (text, ocr_used) — ocr_used is True only for PDFs processed via OCR fallback.
    """
    lower_filename = filename.lower()

    if lower_filename.endswith('.pdf') or content_type == 'application/pdf':
        return extract_text_from_pdf(stream_bytes)
    elif lower_filename.endswith('.docx') or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return extract_text_from_docx(stream_bytes), False
    elif lower_filename.endswith('.doc'):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Legacy binary .doc format is not supported directly. Please convert to .docx or .pdf."
        )
    elif lower_filename.endswith(('.txt', '.md')) or (content_type and content_type.startswith('text/')):
        return extract_text_from_txt(stream_bytes), False
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file format: '{filename}'. Only .pdf, .docx, .txt, and .md files are supported."
        )
